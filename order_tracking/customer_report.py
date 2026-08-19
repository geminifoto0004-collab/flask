"""Customer-facing order report generation (PDF / Word / Excel).

The database keeps status keys in English.  This module is responsible only for
turning those keys and fixed report labels into the requested display language.
Internal workflow notes are never rendered verbatim in customer reports.
"""
from __future__ import annotations

from copy import deepcopy
from datetime import date, datetime, timedelta
from io import BytesIO
import json
import os
import re
import uuid
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from PIL import Image, ImageOps

from .config import (
    UPLOAD_FOLDER,
    CUSTOMER_REPORT_CACHE_DIR,
    CUSTOMER_REPORT_CACHE_HOURS,
    CUSTOMER_REPORT_IMAGE_MAX_EDGE,
    CUSTOMER_REPORT_JPEG_QUALITY,
    CUSTOMER_REPORT_MAX_BYTES,
)
from .status_definitions import get_status_label


REPORT_LANGUAGES = {'zh_cn', 'es'}
REPORT_FORMATS = {'pdf', 'docx', 'xlsx'}
IMAGE_SOURCES = {'both', 'order', 'workflow', 'none'}
IMAGE_COUNTS = {'representative', 'all'}
IMAGE_ORDERS = {'order_first', 'workflow_first', 'newest'}
PDF_ATTACHMENT_MODES = {'pages', 'skip'}

REPORT_TEXT = {
    'zh_cn': {
        'title': '客户订单报告',
        'customer': '客户',
        'order_count': '订单数量',
        'generated_at': '生成日期',
        'order_number': '订单号',
        'order_date': '订单日期',
        'product_name': '产品名称',
        'product_type': '产品类型',
        'product_code': '产品编号',
        'quantity': '数量',
        'status': '目前状态',
        'delivery_date': '交期',
        'shipping_history': '出货记录',
        'no_shipping_history': '暂无出货记录',
        'shipping_pending_confirm': '出货日期待确认',
        'partial_shipment': '部分出货',
        'total_shipment': '全部出货',
        'no_workflow': '尚无流程',
        'images': '图片',
        'image_source_order': '主管参考图',
        'image_source_workflow': '业务员附件图',
        'image_sheet': '订单图片',
        'report_sheet': '订单报告',
        'image_source': '图片来源',
        'image_name': '图片名称',
        'image': '图片',
        'continued': '（续）',
    },
    'es': {
        'title': 'Reporte de pedidos',
        'customer': 'Cliente',
        'order_count': 'Cantidad de pedidos',
        'generated_at': 'Fecha de generación',
        'order_number': 'N.º de pedido',
        'order_date': 'Fecha del pedido',
        'product_name': 'Producto',
        'product_type': 'Tipo de producto',
        'product_code': 'Código de producto',
        'quantity': 'Cantidad',
        'status': 'Estado actual',
        'delivery_date': 'Fecha de entrega',
        'shipping_history': 'Historial de envíos',
        'no_shipping_history': 'Sin registros de envío',
        'shipping_pending_confirm': 'Fecha de envío pendiente de confirmar',
        'partial_shipment': 'Envío parcial',
        'total_shipment': 'Envío total',
        'no_workflow': 'Sin proceso',
        'images': 'Imágenes',
        'image_source_order': 'Imagen de referencia',
        'image_source_workflow': 'Imagen adjunta',
        'image_sheet': 'Imágenes',
        'report_sheet': 'Reporte',
        'image_source': 'Origen',
        'image_name': 'Nombre de imagen',
        'image': 'Imagen',
        'continued': ' (continuación)',
    },
}


def validate_report_options(report_format: str, language: str, image_source: str, image_count: str, image_order: str = 'order_first', pdf_attachment_mode: str = 'pages') -> None:
    if report_format not in REPORT_FORMATS:
        raise ValueError('不支持的报告格式')
    if language not in REPORT_LANGUAGES:
        raise ValueError('不支持的报告语言')
    if image_source not in IMAGE_SOURCES:
        raise ValueError('不支持的图片来源')
    if image_count not in IMAGE_COUNTS:
        raise ValueError('不支持的图片数量选项')
    if image_order not in IMAGE_ORDERS:
        raise ValueError('不支持的图片排序选项')
    if pdf_attachment_mode not in PDF_ATTACHMENT_MODES:
        raise ValueError('不支持的 PDF 附件选项')


def _safe_component(value: str, fallback: str = 'Customer') -> str:
    text = str(value or '').strip() or fallback
    text = re.sub(r'[\\/:*?"<>|\r\n\t]+', '_', text)
    text = re.sub(r'\s+', '_', text).strip('._ ')
    return text[:80] or fallback


def _parse_db_date(value) -> Optional[date]:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value).strip()
    if not text:
        return None
    text = text.split()[0]
    for fmt in ('%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d'):
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            pass
    return None


def format_report_date(value, language: str) -> str:
    d = _parse_db_date(value)
    if not d:
        return str(value or '')
    return d.strftime('%d/%m/%Y') if language == 'es' else d.strftime('%Y-%m-%d')


def format_report_datetime(value, language: str) -> str:
    dt = value if isinstance(value, datetime) else datetime.now()
    return dt.strftime('%d/%m/%Y %H:%M') if language == 'es' else dt.strftime('%Y-%m-%d %H:%M')


def _normalize_status_key(status: str) -> str:
    status = str(status or '').strip()
    if not status:
        return ''
    # get_status_label returns the key itself if unknown.  Try known labels for
    # backwards-compatible databases which may still contain Chinese text.
    from .status_definitions import STATUS_LABELS
    if status in STATUS_LABELS:
        return status
    for key, labels in STATUS_LABELS.items():
        if status in (labels.get('zh_cn'), labels.get('zh_tw'), labels.get('en'), labels.get('es')):
            return key
    return status


def _report_status(status: str, language: str) -> str:
    key = _normalize_status_key(status)
    if not key:
        return REPORT_TEXT[language]['no_workflow']
    return get_status_label(key, language)


def _shipping_records(history: Sequence[dict], workflow_notes: str, order_date, language: str) -> List[str]:
    text = REPORT_TEXT[language]
    records: List[str] = []
    seen = set()
    for row in history or []:
        key = _normalize_status_key(row.get('to_status'))
        if key not in {'PARTIAL_SHIPPED', 'ALL_SHIPPED'}:
            continue
        d = _parse_db_date(row.get('action_date'))
        marker = (key, d.isoformat() if d else str(row.get('action_date') or ''))
        if marker in seen:
            continue
        seen.add(marker)
        label = text['partial_shipment'] if key == 'PARTIAL_SHIPPED' else text['total_shipment']
        records.append(f"{format_report_date(d or row.get('action_date'), language)} - {label}")

    if records:
        return records

    # Legacy transition: parse dates only from the internal convenience note.
    # Never return the original note text to a customer report.
    parsed = _parse_legacy_shipping_note(workflow_notes, order_date)
    for d, kind in parsed:
        label = text['partial_shipment'] if kind == 'partial' else text['total_shipment']
        records.append(f"{format_report_date(d, language)} - {label}")
    if records:
        return records

    note = str(workflow_notes or '')
    if ('出货' in note or '出貨' in note) and note.strip():
        return [text['shipping_pending_confirm']]
    return []


def _parse_legacy_shipping_note(note: str, order_date=None) -> List[Tuple[date, str]]:
    """Extract only high-confidence shipping dates from legacy Chinese notes."""
    note = str(note or '')
    if not note:
        return []
    base = _parse_db_date(order_date) or date.today()
    results: List[Tuple[date, str]] = []
    seen = set()

    # 26.7.18号已出货 / 2026-07-18 出货 / 7-23出货完成
    patterns = [
        re.compile(r'(?:(20\d{2}|\d{2})[.\-/年])?(\d{1,2})[.\-/月](\d{1,2})(?:号|號|日)?[^\n，,。;；]{0,14}(?:出货|出貨)'),
        re.compile(r'(\d{1,2})月(\d{1,2})(?:号|號|日)?[^\n，,。;；]{0,14}(?:出货|出貨)'),
    ]
    for pattern in patterns:
        for m in pattern.finditer(note):
            groups = m.groups()
            if len(groups) == 3:
                y_raw, month_raw, day_raw = groups
                if y_raw:
                    year = int(y_raw)
                    if year < 100:
                        year += 2000
                else:
                    year = base.year
            else:
                year = base.year
                month_raw, day_raw = groups
            try:
                d = date(year, int(month_raw), int(day_raw))
            except ValueError:
                continue
            context = note[max(0, m.start() - 8): min(len(note), m.end() + 16)]
            kind = 'partial' if any(x in context for x in ('部分', '一部份', '一部分', '部份')) else 'total'
            marker = (d, kind)
            if marker not in seen:
                seen.add(marker)
                results.append(marker)
    results.sort(key=lambda x: x[0])
    return results


def load_report_entries(conn, requested_items: Sequence[dict], access_all: bool, current_user_id: Optional[int]) -> List[dict]:
    """Load only rows the current user is allowed to see, preserving UI order."""
    requested = []
    workflow_numbers = []
    order_only_numbers = []
    for item in requested_items or []:
        if not isinstance(item, dict):
            continue
        wf = str(item.get('workflow_number') or '').strip()
        order_no = str(item.get('order_number') or '').strip()
        if not wf and not order_no:
            continue
        key = f'w:{wf}' if wf else f'o:{order_no}'
        if any(x['key'] == key for x in requested):
            continue
        requested.append({'key': key, 'workflow_number': wf, 'order_number': order_no})
        if wf:
            workflow_numbers.append(wf)
        else:
            order_only_numbers.append(order_no)

    if not requested:
        return []

    cursor = conn.cursor()
    workflow_map: Dict[str, dict] = {}
    if workflow_numbers:
        placeholders = ','.join('?' for _ in workflow_numbers)
        sql = f'''
            SELECT w.workflow_number, w.order_number, w.product_name, w.product_code,
                   w.quantity, w.factory, w.production_type, w.expected_delivery_date,
                   w.current_status, w.handler_id, w.notes, w.created_at,
                   o.customer_name, o.order_date, o.visibility, o.status AS order_status,
                   COALESCE(u.real_name, u.display_name, u.username) AS handler_name
            FROM workflows w
            INNER JOIN orders o ON o.order_number = w.order_number
            LEFT JOIN users u ON u.id = w.handler_id
            WHERE w.workflow_number IN ({placeholders})
              AND o.status = 'ACTIVE'
        '''
        params: List[object] = list(workflow_numbers)
        if not access_all:
            sql += " AND w.handler_id = ? AND o.visibility = 'all_sales'"
            params.append(current_user_id)
        cursor.execute(sql, params)
        for row in cursor.fetchall():
            obj = dict(row)
            obj['no_workflow'] = False
            workflow_map[obj['workflow_number']] = obj

    order_map: Dict[str, dict] = {}
    if order_only_numbers:
        placeholders = ','.join('?' for _ in order_only_numbers)
        sql = f'''
            SELECT o.order_number, o.customer_name, o.order_date, o.visibility, o.status AS order_status
            FROM orders o
            WHERE o.order_number IN ({placeholders})
              AND o.status = 'ACTIVE'
              AND NOT EXISTS (SELECT 1 FROM workflows w WHERE w.order_number = o.order_number)
        '''
        params = list(order_only_numbers)
        if not access_all:
            sql += " AND o.visibility = 'all_sales'"
        cursor.execute(sql, params)
        for row in cursor.fetchall():
            obj = dict(row)
            obj.update({
                'workflow_number': '', 'product_name': '', 'product_code': '',
                'quantity': '', 'factory': '', 'production_type': '',
                'expected_delivery_date': '', 'current_status': '', 'handler_id': None,
                'handler_name': '', 'notes': '', 'no_workflow': True,
            })
            order_map[obj['order_number']] = obj

    entries: List[dict] = []
    for req in requested:
        if req['workflow_number']:
            entry = workflow_map.get(req['workflow_number'])
        else:
            entry = order_map.get(req['order_number'])
        if entry:
            entries.append(dict(entry))

    workflow_ids = [e['workflow_number'] for e in entries if e.get('workflow_number')]
    history_map: Dict[str, List[dict]] = {w: [] for w in workflow_ids}
    if workflow_ids:
        placeholders = ','.join('?' for _ in workflow_ids)
        cursor.execute(f'''
            SELECT id, workflow_number, from_status, to_status, action_date, notes, created_at
            FROM workflow_status_history
            WHERE workflow_number IN ({placeholders})
            ORDER BY workflow_number, action_date ASC, created_at ASC, id ASC
        ''', workflow_ids)
        for row in cursor.fetchall():
            obj = dict(row)
            history_map.setdefault(obj['workflow_number'], []).append(obj)

    for entry in entries:
        entry['history'] = history_map.get(entry.get('workflow_number'), [])
    return entries


def _order_file_columns(cursor) -> set:
    cursor.execute('PRAGMA table_info(order_files)')
    return {str(row[1]) for row in cursor.fetchall()}


def _resolve_order_file_path(row: dict) -> Optional[str]:
    raw_path = str(row.get('file_path') or '')
    stored = str(row.get('stored_filename') or '')
    candidates = []
    if raw_path:
        if stored:
            if os.path.isdir(raw_path):
                candidates.append(os.path.join(raw_path, stored))
            candidates.append(os.path.join(UPLOAD_FOLDER, raw_path, stored))
            if os.path.basename(raw_path) == stored:
                candidates.append(raw_path if os.path.isabs(raw_path) else os.path.join(UPLOAD_FOLDER, raw_path))
        else:
            candidates.append(raw_path if os.path.isabs(raw_path) else os.path.join(UPLOAD_FOLDER, raw_path))
    for path in candidates:
        if path and os.path.isfile(path):
            return os.path.normpath(path)
    return None


def _resolve_workflow_file_path(row: dict) -> Optional[str]:
    """Resolve current and legacy workflow attachment paths.

    New rows store file_path as the relative file path itself. Older rows may
    store only the workflow directory; the normal download API already falls
    back to choosing the newest file with the same extension in that folder.
    Customer reports must use the same compatibility rule, otherwise business
    attachments can appear in the UI but disappear from exported reports.
    """
    raw_path = str(row.get('file_path') or '').strip()
    if not raw_path:
        return None

    original_name = str(row.get('display_name') or '').strip()
    if os.path.isabs(raw_path):
        candidates = [raw_path]
    else:
        candidates = [os.path.join(UPLOAD_FOLDER, raw_path)]

    for path in candidates:
        if os.path.isfile(path):
            return os.path.normpath(path)

        # Legacy data may point at a directory rather than the physical file.
        dir_path = path if os.path.isdir(path) else None
        if not dir_path:
            candidate_dir = path.rstrip('/\\')
            if os.path.isdir(candidate_dir):
                dir_path = candidate_dir
        if not dir_path:
            continue

        ext = os.path.splitext(original_name)[1].lower()
        try:
            files = [
                name for name in os.listdir(dir_path)
                if os.path.isfile(os.path.join(dir_path, name))
                and (not ext or name.lower().endswith(ext))
            ]
        except OSError:
            files = []
        if files:
            files.sort(key=lambda name: os.path.getmtime(os.path.join(dir_path, name)), reverse=True)
            return os.path.normpath(os.path.join(dir_path, files[0]))

    return None


def attach_image_metadata(conn, entries: List[dict], image_source: str, image_count: str, image_order: str = 'order_first', pdf_attachment_mode: str = 'pages') -> None:
    for e in entries:
        e['images'] = []
        e['_source_image_meta'] = {'order': [], 'workflow': []}
        e['_pdf_attachment_count'] = 0
    if image_source == 'none' or not entries:
        return

    cursor = conn.cursor()
    order_numbers = sorted({e.get('order_number') for e in entries if e.get('order_number')})
    workflow_numbers = sorted({e.get('workflow_number') for e in entries if e.get('workflow_number')})
    order_images: Dict[str, List[dict]] = {x: [] for x in order_numbers}
    workflow_images: Dict[str, List[dict]] = {x: [] for x in workflow_numbers}

    if image_source in {'both', 'order'} and order_numbers:
        cols = _order_file_columns(cursor)
        filename_expr = 'original_filename' if 'original_filename' in cols else ('file_name' if 'file_name' in cols else "''")
        stored_expr = 'stored_filename' if 'stored_filename' in cols else "''"
        mime_expr = 'mime_type' if 'mime_type' in cols else ('file_type' if 'file_type' in cols else "''")
        uploaded_expr = 'uploaded_at' if 'uploaded_at' in cols else 'NULL'
        placeholders = ','.join('?' for _ in order_numbers)
        cursor.execute(f'''
            SELECT id, order_number, {filename_expr} AS display_name,
                   {stored_expr} AS stored_filename, file_path,
                   file_size, {mime_expr} AS mime_type, {uploaded_expr} AS uploaded_at
            FROM order_files
            WHERE order_number IN ({placeholders})
            ORDER BY uploaded_at DESC, id DESC
        ''', order_numbers)
        for row in cursor.fetchall():
            obj = dict(row)
            path = _resolve_order_file_path(obj)
            if path and _is_report_visual_file(path, obj.get('mime_type')):
                obj.update({'path': path, 'source': 'order', 'media_type': 'pdf' if _is_pdf_file(path, obj.get('mime_type')) else 'image'})
                order_images.setdefault(obj['order_number'], []).append(obj)

    if image_source in {'both', 'workflow'} and workflow_numbers:
        placeholders = ','.join('?' for _ in workflow_numbers)
        cursor.execute(f'''
            SELECT id, workflow_number, file_name AS display_name, file_path,
                   file_size, file_type AS mime_type, uploaded_at
            FROM workflow_files
            WHERE workflow_number IN ({placeholders}) AND is_deleted = 0
            ORDER BY uploaded_at DESC, id DESC
        ''', workflow_numbers)
        for row in cursor.fetchall():
            obj = dict(row)
            path = _resolve_workflow_file_path(obj)
            if path and _is_report_visual_file(path, obj.get('mime_type')):
                obj.update({'path': path, 'source': 'workflow', 'media_type': 'pdf' if _is_pdf_file(path, obj.get('mime_type')) else 'image'})
                workflow_images.setdefault(obj['workflow_number'], []).append(obj)

    for entry in entries:
        order_list = list(order_images.get(entry.get('order_number'), []))
        workflow_list = list(workflow_images.get(entry.get('workflow_number'), []))
        # 各来源内部都维持最新上传在前；合并时依使用者选择排序。
        order_list.sort(key=lambda x: str(x.get('uploaded_at') or ''), reverse=True)
        workflow_list.sort(key=lambda x: str(x.get('uploaded_at') or ''), reverse=True)

        # Always count selected PDF attachments so the UI can ask the user what to do
        # even when the current choice is "skip". The original PDF is never modified.
        selected_for_pdf_count = []
        if image_source in {'both', 'order'}:
            selected_for_pdf_count.extend(order_list)
        if image_source in {'both', 'workflow'}:
            selected_for_pdf_count.extend(workflow_list)
        entry['_pdf_attachment_count'] = sum(
            1 for meta in selected_for_pdf_count
            if _is_pdf_file(meta.get('path') or '', meta.get('mime_type'))
        )

        if pdf_attachment_mode == 'skip':
            order_list = [x for x in order_list if not _is_pdf_file(x.get('path') or '', x.get('mime_type'))]
            workflow_list = [x for x in workflow_list if not _is_pdf_file(x.get('path') or '', x.get('mime_type'))]

        if image_source == 'order':
            images = order_list
        elif image_source == 'workflow':
            images = workflow_list
        elif image_order == 'workflow_first':
            images = workflow_list + order_list
        elif image_order == 'newest':
            images = order_list + workflow_list
            images.sort(key=lambda x: str(x.get('uploaded_at') or ''), reverse=True)
        else:
            # 默认：主管参考图优先，再放业务员附件图。
            images = order_list + workflow_list
        # De-duplicate the same physical path. Keep source-specific pools too: Excel
        # uses one representative image from EACH selected source on its main sheet,
        # even when the report option is "one representative image".
        def _dedupe_source(source_items):
            out, source_seen = [], set()
            for source_img in source_items:
                key = os.path.normcase(os.path.abspath(source_img['path']))
                if key in source_seen:
                    continue
                source_seen.add(key)
                out.append(source_img)
            return out

        order_list = _dedupe_source(order_list)
        workflow_list = _dedupe_source(workflow_list)
        entry['_source_image_meta'] = {'order': order_list, 'workflow': workflow_list}

        deduped = []
        seen = set()
        for img in images:
            key = os.path.normcase(os.path.abspath(img['path']))
            if key in seen:
                continue
            seen.add(key)
            deduped.append(img)
        entry['images'] = deduped[:1] if image_count == 'representative' else deduped


def _is_image_file(path: str, mime_type: str = '') -> bool:
    mime = str(mime_type or '').lower()
    ext = os.path.splitext(path)[1].lower()
    # 这里只做轻量判断，避免 Modal 每次预估都把所有图片逐张打开验证。
    # 真正导出时 _compress_image() 会负责最终解码验证，坏图会自动跳过。
    return mime.startswith('image/') or ext in {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.gif', '.tif', '.tiff'}


def _is_pdf_file(path: str, mime_type: str = '') -> bool:
    mime = str(mime_type or '').lower().split(';', 1)[0].strip()
    ext = os.path.splitext(path)[1].lower()
    return mime == 'application/pdf' or ext == '.pdf'


def _is_report_visual_file(path: str, mime_type: str = '') -> bool:
    """Images and PDF attachments are both visual sources for customer reports."""
    return _is_image_file(path, mime_type) or _is_pdf_file(path, mime_type)


def _pdf_attachment_pages(path: str, first_page_only: bool = False) -> List[dict]:
    """Render a PDF attachment into JPEG page images for the final report.

    The original PDF is never modified.  Rendering happens only inside the report
    worker.  Pages are compressed with the same limits as normal uploaded images.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise RuntimeError(
            '检测到 PDF 附件，但当前 Flask Python 环境未安装 PyMuPDF。'
            '请执行: python -m pip install PyMuPDF'
        ) from exc

    rendered: List[dict] = []
    try:
        with fitz.open(path) as doc:
            total_pages = len(doc)
            limit = min(total_pages, 1) if first_page_only else total_pages
            for page_index in range(limit):
                page = doc.load_page(page_index)
                # 120 DPI is enough for the existing ~1200px customer-report image target.
                pix = page.get_pixmap(dpi=120, colorspace=fitz.csRGB, alpha=False)
                img = Image.frombytes('RGB', (pix.width, pix.height), pix.samples)
                img.thumbnail(
                    (CUSTOMER_REPORT_IMAGE_MAX_EDGE, CUSTOMER_REPORT_IMAGE_MAX_EDGE),
                    Image.Resampling.LANCZOS,
                )
                out = BytesIO()
                img.save(
                    out,
                    format='JPEG',
                    quality=CUSTOMER_REPORT_JPEG_QUALITY,
                    optimize=False,
                )
                payload = out.getvalue()
                rendered.append({
                    'bytes': payload,
                    'width': img.width,
                    'height': img.height,
                    'file_size': len(payload),
                    'pdf_page_number': page_index + 1,
                    'pdf_page_count': total_pages,
                })
    except RuntimeError:
        raise
    except Exception as exc:
        print(f'[WARN] customer report PDF attachment skipped: {path}: {exc}')
        return []
    return rendered


def _compress_image(path: str) -> Optional[dict]:
    try:
        with Image.open(path) as original:
            img = ImageOps.exif_transpose(original)
            if getattr(img, 'is_animated', False):
                img.seek(0)
            if img.mode in ('RGBA', 'LA'):
                background = Image.new('RGB', img.size, 'white')
                alpha = img.getchannel('A') if 'A' in img.getbands() else None
                background.paste(img.convert('RGB'), mask=alpha)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            img.thumbnail((CUSTOMER_REPORT_IMAGE_MAX_EDGE, CUSTOMER_REPORT_IMAGE_MAX_EDGE), Image.Resampling.LANCZOS)
            out = BytesIO()
            # optimize=True 会为每张图额外做一轮昂贵扫描；客户报告更重视生成速度。
            img.save(out, format='JPEG', quality=CUSTOMER_REPORT_JPEG_QUALITY, optimize=False)
            return {'bytes': out.getvalue(), 'width': img.width, 'height': img.height}
    except Exception:
        return None


def prepare_entries(conn, entries: List[dict], image_source: str, image_count: str, language: str, image_order: str = 'order_first', pdf_attachment_mode: str = 'pages') -> List[dict]:
    attach_image_metadata(conn, entries, image_source, image_count, image_order, pdf_attachment_mode)
    prepared = []
    image_cache: Dict[str, Optional[dict]] = {}
    pdf_cache: Dict[Tuple[str, bool], List[dict]] = {}
    for entry in entries:
        item = dict(entry)
        item['image_order'] = image_order
        item['pdf_attachment_mode'] = pdf_attachment_mode
        item['status_label'] = _report_status(item.get('current_status'), language)
        item['shipping_records'] = _shipping_records(
            item.get('history') or [], item.get('notes') or '', item.get('order_date'), language
        )
        images = []
        for meta in item.get('images') or []:
            path = meta['path']
            is_pdf = _is_pdf_file(path, meta.get('mime_type'))
            if is_pdf:
                cache_key = (path, image_count == 'representative')
                if cache_key not in pdf_cache:
                    pdf_cache[cache_key] = _pdf_attachment_pages(
                        path,
                        first_page_only=(image_count == 'representative'),
                    )
                base_name = str(meta.get('display_name') or os.path.basename(path) or 'PDF')
                for page_data in pdf_cache[cache_key]:
                    obj = dict(meta)
                    obj.update(page_data)
                    obj['media_type'] = 'pdf_page'
                    obj['original_pdf_path'] = path
                    obj['display_name'] = (
                        f"{base_name} [PDF {page_data['pdf_page_number']}/{page_data['pdf_page_count']}]"
                    )
                    images.append(obj)
                continue

            if path not in image_cache:
                image_cache[path] = _compress_image(path)
            compressed = image_cache[path]
            if not compressed:
                continue
            obj = dict(meta)
            obj.update(compressed)
            images.append(obj)
        item['images'] = images

        # Excel main sheet needs an independent representative from each source.
        # This is intentionally separate from item['images'], because PDF/Word's
        # "representative" option still means one image for the order as before.
        # If that source's representative attachment is a PDF, use its first page.
        excel_source_images = {}
        source_meta = item.get('_source_image_meta') or {}
        for source_key in ('order', 'workflow'):
            source_items = source_meta.get(source_key) or []
            for meta in source_items:
                path = meta.get('path')
                if not path:
                    continue
                if _is_pdf_file(path, meta.get('mime_type')):
                    cache_key = (path, True)
                    if cache_key not in pdf_cache:
                        pdf_cache[cache_key] = _pdf_attachment_pages(path, first_page_only=True)
                    if not pdf_cache[cache_key]:
                        continue
                    page_data = pdf_cache[cache_key][0]
                    base_name = str(meta.get('display_name') or os.path.basename(path) or 'PDF')
                    obj = dict(meta)
                    obj.update(page_data)
                    obj['media_type'] = 'pdf_page'
                    obj['original_pdf_path'] = path
                    obj['display_name'] = (
                        f"{base_name} [PDF {page_data['pdf_page_number']}/{page_data['pdf_page_count']}]"
                    )
                    excel_source_images[source_key] = obj
                    break

                if path not in image_cache:
                    image_cache[path] = _compress_image(path)
                compressed = image_cache[path]
                if not compressed:
                    continue
                obj = dict(meta)
                obj.update(compressed)
                excel_source_images[source_key] = obj
                break
        item['excel_source_images'] = excel_source_images
        item.pop('_source_image_meta', None)
        item.pop('_pdf_attachment_count', None)
        prepared.append(item)
    return prepared


def _group_entries_by_customer(entries: Sequence[dict]) -> List[Tuple[str, List[dict]]]:
    """Preserve UI order while grouping reports by customer."""
    groups: List[Tuple[str, List[dict]]] = []
    index: Dict[str, int] = {}
    for entry in entries or []:
        customer = str(entry.get('customer_name') or '').strip() or 'Customer'
        key = customer.casefold()
        if key not in index:
            index[key] = len(groups)
            groups.append((customer, []))
        groups[index[key]][1].append(entry)
    return groups


def estimate_report(entries: Sequence[dict], report_format: str) -> dict:
    estimated = _estimate_bytes(entries)
    customer_groups = _group_entries_by_customer(entries)
    customer_count = max(1, len(customer_groups)) if entries else 0

    if report_format == 'xlsx':
        return {
            'estimated_bytes': estimated,
            'estimated_parts': 1,
            'estimated_files': 1,
            'customer_count': customer_count,
        }

    if report_format == 'pdf' and customer_groups:
        # PDF 是给客户的文件：不同客户永远分开生成，不能混在同一份文件。
        estimated_files = 0
        for _, group in customer_groups:
            group_bytes = _estimate_bytes(group)
            estimated_files += max(1, (group_bytes + CUSTOMER_REPORT_MAX_BYTES - 1) // CUSTOMER_REPORT_MAX_BYTES)
        return {
            'estimated_bytes': estimated,
            'estimated_parts': int(estimated_files),
            'estimated_files': int(estimated_files),
            'customer_count': len(customer_groups),
        }

    parts = max(1, (estimated + CUSTOMER_REPORT_MAX_BYTES - 1) // CUSTOMER_REPORT_MAX_BYTES)
    return {
        'estimated_bytes': estimated,
        'estimated_parts': int(parts),
        'estimated_files': int(parts),
        'customer_count': customer_count,
    }


def _estimate_bytes(entries: Sequence[dict]) -> int:
    base = 90_000 + len(entries) * 5_000
    images = 0
    for e in entries:
        for img in e.get('images') or []:
            raw = int(img.get('file_size') or 0)
            # Typical 1000-1200px JPEG after compression. Keep estimate conservative.
            images += min(max(int(raw * 0.40), 80_000), 550_000) if raw else 220_000
    return base + images




def build_report_files(entries: List[dict], report_format: str, language: str) -> List[Tuple[str, str, bytes]]:
    validate_report_options(report_format, language, 'none', 'representative')
    if report_format == 'xlsx':
        data = _render_excel(entries, language)
        name = _base_filename(entries, language, 'xlsx')
        return [(name, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', data)]

    # PDF 是客户文件：如果当前筛选里有两个或更多客户，必须按客户分开。
    if report_format == 'pdf':
        mimetype = 'application/pdf'
        files: List[Tuple[str, str, bytes]] = []
        customer_groups = _group_entries_by_customer(entries)
        for _, group_entries in customer_groups:
            rendered_parts = _render_with_size_split(group_entries, language, _render_pdf)
            base = _base_filename(group_entries, language, 'pdf', include_ext=False)
            multi = len(rendered_parts) > 1
            for idx, data in enumerate(rendered_parts, 1):
                filename = f'{base}_Part{idx}.pdf' if multi else f'{base}.pdf'
                files.append((filename, mimetype, data))
        return files


    renderer = _render_docx
    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    rendered_parts = _render_with_size_split(entries, language, renderer)
    base = _base_filename(entries, language, 'docx', include_ext=False)
    files = []
    multi = len(rendered_parts) > 1
    for idx, data in enumerate(rendered_parts, 1):
        filename = f'{base}_Part{idx}.docx' if multi else f'{base}.docx'
        files.append((filename, mimetype, data))
    return files


def _base_filename(entries: Sequence[dict], language: str, ext: str, include_ext: bool = True) -> str:
    customers = []
    for e in entries:
        name = str(e.get('customer_name') or '').strip()
        if name and name not in customers:
            customers.append(name)
    customer = _safe_component(customers[0], 'Customer') if len(customers) == 1 else ('Clientes' if language == 'es' else '客户')
    if len(entries) == 1:
        entry = entries[0] or {}
        order_ref = _safe_component(str(entry.get('workflow_number') or entry.get('order_number') or '').strip(), 'Pedido')
        label = 'ReportePedido' if language == 'es' else '订单报告'
        base = f'{customer}_{order_ref}_{label}_{date.today().strftime("%Y%m%d")}'
    else:
        label = 'ReportePedidos' if language == 'es' else '订单报告'
        base = f'{customer}_{label}_{date.today().strftime("%Y%m%d")}'
    return f'{base}.{ext}' if include_ext else base


def _render_with_size_split(entries: List[dict], language: str, renderer) -> List[bytes]:
    if not entries:
        return [renderer([], language)]

    def recurse(items: List[dict]) -> List[bytes]:
        data = renderer(items, language)
        if len(data) <= CUSTOMER_REPORT_MAX_BYTES:
            return [data]
        if len(items) > 1:
            mid = len(items) // 2
            return recurse(items[:mid]) + recurse(items[mid:])
        item = items[0]
        images = item.get('images') or []
        if len(images) > 1:
            mid = len(images) // 2
            first = deepcopy(item)
            second = deepcopy(item)
            first['images'] = images[:mid]
            second['images'] = images[mid:]
            second['continued'] = True
            return recurse([first]) + recurse([second])
        return [data]

    return recurse(entries)


def _render_pdf(entries: List[dict], language: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        Flowable, Image as RLImage, KeepTogether, PageBreak, Paragraph,
        SimpleDocTemplate, Spacer, Table, TableStyle
    )

    out = BytesIO()
    doc = SimpleDocTemplate(
        out, pagesize=A4,
        leftMargin=13 * mm, rightMargin=13 * mm,
        topMargin=15 * mm, bottomMargin=14 * mm,
        title=REPORT_TEXT[language]['title'], author='order_tracking'
    )
    styles = getSampleStyleSheet()
    cjk_font = 'STSong-Light'
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(cjk_font))
    except Exception:
        pass

    if language == 'zh_cn':
        font_name = cjk_font
        bold_font = cjk_font
    else:
        font_name = 'Helvetica'
        bold_font = 'Helvetica-Bold'

    # Match the main tracking page: crisp dark text, warm white cards and restrained brand red.
    ink = colors.HexColor('#171717')
    ink_2 = colors.HexColor('#343434')
    muted = colors.HexColor('#667085')
    line = colors.HexColor('#F0D7DB')
    soft = colors.HexColor('#FFF8F8')
    soft_2 = colors.HexColor('#FFF0F0')
    accent = colors.HexColor('#FF2442')
    accent_dark = colors.HexColor('#D91E39')
    sand = colors.HexColor('#FFD4DA')
    sand_soft = colors.HexColor('#FFF7F8')
    danger = colors.HexColor('#C91831')

    title_style = ParagraphStyle(
        'CRTitleV7', parent=styles['Title'], fontName=bold_font,
        fontSize=20, leading=23, textColor=ink, alignment=TA_LEFT, spaceAfter=0
    )
    customer_style = ParagraphStyle(
        'CRCustomerV7', parent=styles['BodyText'], fontName=bold_font,
        fontSize=10.2, leading=12.3, textColor=ink, alignment=TA_LEFT
    )
    eyebrow_style = ParagraphStyle(
        'CREyebrowV7', parent=styles['BodyText'], fontName='Helvetica-Bold',
        fontSize=7.2, leading=8.5, textColor=accent_dark, alignment=TA_LEFT,
        spaceAfter=0
    )
    meta_style = ParagraphStyle(
        'CRMetaV7', parent=styles['BodyText'], fontName=font_name,
        fontSize=8.0, leading=10.2, textColor=muted, alignment=TA_RIGHT
    )
    overview_title_style = ParagraphStyle(
        'CROverviewTitle', parent=styles['Heading2'], fontName=bold_font,
        fontSize=11.5, leading=14, textColor=ink, spaceAfter=0
    )
    order_num_style = ParagraphStyle(
        'CROrderNumV7', parent=styles['Heading2'], fontName=bold_font,
        fontSize=15.2, leading=18.5, textColor=ink, spaceBefore=0, spaceAfter=0
    )
    order_status_style = ParagraphStyle(
        'CROrderStatusV7', parent=styles['BodyText'], fontName=bold_font,
        fontSize=12.2, leading=14.2, textColor=accent_dark, alignment=TA_RIGHT
    )
    hero_label_style = ParagraphStyle(
        'CRHeroLabelV7', parent=styles['BodyText'], fontName=font_name,
        fontSize=7.3, leading=8.5, textColor=muted, spaceAfter=1.4
    )
    hero_value_style = ParagraphStyle(
        'CRHeroValueV7', parent=styles['BodyText'], fontName=bold_font,
        fontSize=13.2, leading=15.2, textColor=ink, spaceAfter=0
    )
    hero_sub_style = ParagraphStyle(
        'CRHeroSubV7', parent=styles['BodyText'], fontName=font_name,
        fontSize=7.7, leading=9.3, textColor=muted, spaceAfter=0
    )
    key_label_style = ParagraphStyle(
        'CRKeyLabelV7', parent=styles['BodyText'], fontName=font_name,
        fontSize=7.2, leading=8.4, textColor=muted, spaceAfter=1.5
    )
    key_value_style = ParagraphStyle(
        'CRKeyValueV7', parent=styles['BodyText'], fontName=bold_font,
        fontSize=10.0, leading=11.8, textColor=ink
    )
    section_style = ParagraphStyle(
        'CRSectionV7', parent=styles['Heading3'], fontName=bold_font,
        fontSize=9.5, leading=11.5, textColor=ink, spaceBefore=0, spaceAfter=0
    )
    body_style = ParagraphStyle(
        'CRBodyV7', parent=styles['BodyText'], fontName=font_name,
        fontSize=8.7, leading=11.2, textColor=ink
    )
    small_style = ParagraphStyle(
        'CRSmallV7', parent=styles['BodyText'], fontName=font_name,
        fontSize=7.4, leading=9.1, textColor=muted
    )
    image_caption_style = ParagraphStyle(
        'CRImageCaptionV7', parent=body_style, fontSize=6.8, leading=8.3,
        textColor=muted, alignment=TA_CENTER
    )
    image_source_style = ParagraphStyle(
        'CRImageSourceV7', parent=body_style, fontName=bold_font, fontSize=6.9,
        leading=8.3, textColor=ink_2, alignment=TA_LEFT
    )

    text = REPORT_TEXT[language]
    customers = _distinct_customers(entries)
    customer_name = customers[0] if len(customers) == 1 else ''
    story = []

    if language == 'es':
        labels = {
            'overview': 'Resumen de pedidos',
            'overview_note': 'Vista rápida del avance, entrega y estado de cada pedido.',
            'pattern': 'Producto / diseño',
            'delivery': 'Entrega prevista',
            'current_status_card': 'Estado actual',
            'quantity': 'Cantidad',
            'current_progress': 'Progreso actual',
            'order_date': 'Pedido',
            'code': 'Código',
            'type': 'Tipo',
            'pending_date': 'Por confirmar',
            'confirmed': 'Confirmado',
            'pending_confirm': 'Pendiente de confirmación',
            'revising': 'En revisión',
            'preparing': 'En preparación',
            'cancelled': 'Cancelado',
            'shipping': 'Envío',
            'stage_names': ['Pedido', 'Diseño', 'Muestra', 'Producción', 'Envío'],
            'detail': 'Detalle del pedido',
        }
    else:
        labels = {
            'overview': '订单总览',
            'overview_note': '一眼查看每张订单目前走到哪个阶段。',
            'pattern': '花型 / 产品',
            'delivery': '预计交期',
            'current_status_card': '当前状态',
            'quantity': '数量',
            'current_progress': '当前进度',
            'order_date': '下单日期',
            'code': '产品编号',
            'type': '产品类型',
            'pending_date': '待确认',
            'confirmed': '已确认',
            'pending_confirm': '待确认',
            'revising': '修改中',
            'preparing': '准备中',
            'cancelled': '已取消',
            'shipping': '出货',
            'stage_names': ['订单', '图稿', '打样', '生产', '出货'],
            'detail': '订单详情',
        }

    def status_key(entry):
        return _normalize_status_key(entry.get('current_status'))

    def confirmation_text(key: str) -> str:
        if key == 'CANCELLED':
            return labels['cancelled']
        if key in {'QUOTE_CONFIRMING', 'DRAFT_CONFIRMING', 'SAMPLE_CONFIRMING'}:
            if language == 'es':
                detail = {
                    'QUOTE_CONFIRMING': 'Cotización pendiente',
                    'DRAFT_CONFIRMING': 'Diseño pendiente',
                    'SAMPLE_CONFIRMING': 'Muestra pendiente',
                }
            else:
                detail = {
                    'QUOTE_CONFIRMING': '报价待确认',
                    'DRAFT_CONFIRMING': '图稿待确认',
                    'SAMPLE_CONFIRMING': '样品待确认',
                }
            return detail[key]
        if key in {'DRAFT_REVISING', 'SAMPLE_REVISING'}:
            return labels['revising']
        if key in {'PENDING_PRODUCTION', 'PRODUCING', 'PRODUCTION_DONE', 'PARTIAL_SHIPPED', 'ALL_SHIPPED', 'COMPLETED'}:
            return labels['confirmed']
        if key in {'PENDING_SAMPLE', 'SAMPLING'}:
            return 'Diseño confirmado' if language == 'es' else '图稿已确认'
        return labels['preparing']

    def stage_index(key: str) -> int:
        if key in {'NEW_ORDER', 'QUOTE_CONFIRMING'}:
            return 0
        if key in {'DRAFT_MAKING', 'DRAFT_CONFIRMING', 'DRAFT_REVISING'}:
            return 1
        if key in {'PENDING_SAMPLE', 'SAMPLING', 'SAMPLE_CONFIRMING', 'SAMPLE_REVISING'}:
            return 2
        if key in {'PENDING_PRODUCTION', 'PRODUCING', 'PRODUCTION_DONE'}:
            return 3
        if key in {'PARTIAL_SHIPPED', 'ALL_SHIPPED', 'COMPLETED'}:
            return 4
        return 0

    def delivery_text(entry):
        raw = entry.get('expected_delivery_date')
        return format_report_date(raw, language) if _parse_db_date(raw) else labels['pending_date']

    def short_report_date(value):
        d = _parse_db_date(value)
        if not d:
            return ''
        return d.strftime('%d/%m') if language == 'es' else d.strftime('%m/%d')

    def overview_stage_index(entry):
        key = status_key(entry)
        if key != 'CANCELLED':
            return stage_index(key)
        # If an order was cancelled after work had already progressed, keep the
        # last real stage visible instead of making the timeline jump back to 0.
        for row in reversed(entry.get('history') or []):
            hist_key = _normalize_status_key(row.get('to_status'))
            if hist_key and hist_key != 'CANCELLED':
                return stage_index(hist_key)
        return 0

    def overview_stage_dates(entry):
        dates = [''] * 5
        dates[0] = short_report_date(entry.get('order_date'))
        for row in entry.get('history') or []:
            hist_key = _normalize_status_key(row.get('to_status'))
            if not hist_key or hist_key == 'CANCELLED':
                continue
            idx = stage_index(hist_key)
            if idx <= 0 or dates[idx]:
                continue
            dates[idx] = short_report_date(row.get('action_date') or row.get('created_at'))
        return dates

    def product_primary(entry):
        return str(entry.get('product_name') or entry.get('production_type') or '-').strip() or '-'

    class MiniStageTimeline(Flowable):
        def __init__(self, current_idx: int, stage_dates=None, width=120 * mm, height=16 * mm, cancelled=False):
            Flowable.__init__(self)
            self.current_idx = max(0, min(4, current_idx))
            self.stage_dates = list(stage_dates or [''] * 5)[:5]
            while len(self.stage_dates) < 5:
                self.stage_dates.append('')
            self.width = width
            self.height = height
            self.cancelled = bool(cancelled)

        def draw(self):
            c = self.canv
            c.saveState()
            left = 6 * mm
            right = self.width - 6 * mm
            y = 7.1 * mm
            step = (right - left) / 4.0

            c.setLineWidth(1.0)
            c.setStrokeColor(line)
            c.line(left, y, right, y)
            if self.current_idx > 0 and not self.cancelled:
                c.setStrokeColor(accent)
                c.setLineWidth(1.25)
                c.line(left, y, left + step * self.current_idx, y)

            label_size = 5.9 if language == 'es' else 6.2
            date_size = 5.8
            for i, label in enumerate(labels['stage_names']):
                x = left + step * i
                reached = i <= self.current_idx and not self.cancelled
                if i < self.current_idx and not self.cancelled:
                    c.setFillColor(accent)
                    c.setStrokeColor(accent)
                    c.circle(x, y, 1.75 * mm, stroke=0, fill=1)
                elif i == self.current_idx and not self.cancelled:
                    c.setFillColor(colors.white)
                    c.setStrokeColor(accent)
                    c.setLineWidth(1.25)
                    c.circle(x, y, 2.05 * mm, stroke=1, fill=1)
                    c.setFillColor(accent)
                    c.circle(x, y, .78 * mm, stroke=0, fill=1)
                else:
                    c.setFillColor(colors.white)
                    c.setStrokeColor(line)
                    c.setLineWidth(.9)
                    c.circle(x, y, 1.7 * mm, stroke=1, fill=1)

                date_label = self.stage_dates[i] if reached else ''
                if date_label:
                    c.setFillColor(accent_dark if i == self.current_idx else muted)
                    c.setFont(font_name, date_size)
                    tw = pdfmetrics.stringWidth(date_label, font_name, date_size)
                    c.drawString(x - tw / 2, 11.2 * mm, date_label)

                c.setFillColor(ink if reached else muted)
                c.setFont(font_name, label_size)
                tw = pdfmetrics.stringWidth(label, font_name, label_size)
                c.drawString(x - tw / 2, .6 * mm, label)
            c.restoreState()

    class TextileMotif(Flowable):
        def __init__(self, width=25 * mm, height=17 * mm):
            Flowable.__init__(self)
            self.width = width
            self.height = height
        def draw(self):
            c = self.canv
            c.saveState()
            c.setStrokeColor(line)
            c.setLineWidth(.55)
            w, h = self.width, self.height
            # small woven / repeat-pattern motif: decorative, vector-only, no external asset
            for i in range(5):
                x = 2 * mm + i * 4.7 * mm
                c.line(x, 2 * mm, x + 8 * mm, h - 2 * mm)
                c.line(x, h - 2 * mm, x + 8 * mm, 2 * mm)
            c.setStrokeColor(accent)
            c.setLineWidth(1.25)
            c.circle(w - 5.8 * mm, h / 2, 3.2 * mm, stroke=1, fill=0)
            c.setFillColor(sand)
            c.circle(w - 5.8 * mm, h / 2, 1.25 * mm, stroke=0, fill=1)
            c.restoreState()

    class StageTimeline(Flowable):
        def __init__(self, current_idx: int, width=178 * mm, height=18 * mm):
            Flowable.__init__(self)
            self.current_idx = max(0, min(4, current_idx))
            self.width = width
            self.height = height
        def draw(self):
            c = self.canv
            c.saveState()
            left = 8 * mm
            right = self.width - 8 * mm
            y = 8.0 * mm
            step = (right - left) / 4.0
            c.setLineWidth(1.2)
            c.setStrokeColor(line)
            c.line(left, y, right, y)
            if self.current_idx > 0:
                c.setStrokeColor(accent)
                c.line(left, y, left + step * self.current_idx, y)
            for i, label in enumerate(labels['stage_names']):
                x = left + step * i
                if i < self.current_idx:
                    c.setFillColor(accent)
                    c.circle(x, y, 2.25 * mm, stroke=0, fill=1)
                    c.setStrokeColor(colors.white)
                    c.setLineWidth(1.0)
                    c.line(x - .8 * mm, y, x - .1 * mm, y - .8 * mm)
                    c.line(x - .1 * mm, y - .8 * mm, x + 1.1 * mm, y + .9 * mm)
                elif i == self.current_idx:
                    c.setFillColor(colors.white)
                    c.setStrokeColor(accent)
                    c.setLineWidth(1.6)
                    c.circle(x, y, 2.7 * mm, stroke=1, fill=1)
                    c.setFillColor(accent)
                    c.circle(x, y, 1.05 * mm, stroke=0, fill=1)
                else:
                    c.setFillColor(colors.white)
                    c.setStrokeColor(line)
                    c.setLineWidth(1.0)
                    c.circle(x, y, 2.15 * mm, stroke=1, fill=1)
                c.setFillColor(ink if i <= self.current_idx else muted)
                c.setFont(font_name, 6.5 if language == 'es' else 6.8)
                tw = pdfmetrics.stringWidth(label, font_name, 6.5 if language == 'es' else 6.8)
                c.drawString(x - tw / 2, 1.2 * mm, label)
            c.restoreState()

    # ---------- Report identity / customer summary ----------
    customer_markup = _pdf_mixed_markup(customer_name or '-', cjk_font)
    header_left = [
        Paragraph('ORDER STATUS REPORT' if language == 'es' else _pdf_mixed_markup('订单状态报告', cjk_font), eyebrow_style),
        Spacer(1, 1.0 * mm),
        Paragraph(text['title'], title_style),
        Spacer(1, .8 * mm),
        Paragraph(customer_markup, customer_style),
    ]
    header_meta = Paragraph(
        f"{text['order_count']}: <b>{len(entries)}</b><br/>{text['generated_at']}: {format_report_datetime(datetime.now(), language)}",
        meta_style,
    )
    header = Table([[header_left, header_meta, TextileMotif()]], colWidths=[108 * mm, 49 * mm, 27 * mm])
    header.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(header)
    story.append(Spacer(1, 2.3 * mm))
    accent_line = Table([['']], colWidths=[184 * mm], rowHeights=[0.7 * mm])
    accent_line.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), accent)]))
    story.append(accent_line)
    story.append(Spacer(1, 2.2 * mm))

    # ---------- First section: a customer-centric overview ----------
    if len(entries) > 1:
        overview_head = Table([[
            Paragraph(_pdf_mixed_markup(labels['overview'], cjk_font), overview_title_style),
            Paragraph(_pdf_mixed_markup(labels['overview_note'], cjk_font), small_style),
        ]], colWidths=[43 * mm, 141 * mm])
        overview_head.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'BASELINE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0), ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(overview_head)
        story.append(Spacer(1, 2.0 * mm))

        overview_order_style = ParagraphStyle(
            'CROverviewOrderV2', parent=body_style, fontName=bold_font,
            fontSize=10.6, leading=12.4, textColor=ink, spaceAfter=1.0
        )
        overview_meta_style = ParagraphStyle(
            'CROverviewMetaV2', parent=small_style, fontName=font_name,
            fontSize=7.0, leading=8.5, textColor=muted, spaceAfter=.5
        )
        overview_status_style = ParagraphStyle(
            'CROverviewStatusV2', parent=small_style, fontName=bold_font,
            fontSize=10.6, leading=12.3, textColor=accent_dark, spaceAfter=0
        )

        for ov_idx, e in enumerate(entries):
            key = status_key(e)
            order_display = e.get('workflow_number') or e.get('order_number') or '-'
            status_label = e.get('status_label') or _report_status(e.get('current_status'), language)
            product = product_primary(e)
            left_bits = [
                Paragraph(f"{text['order_number']} <b>{_pdf_mixed_markup(order_display, cjk_font)}</b>", overview_order_style),
                Paragraph(_pdf_mixed_markup(product, cjk_font), overview_meta_style),
            ]
            # 客户真正关心的是当前状态。没有明确交期时，不再显示“待确认 / Por confirmar”。
            if _parse_db_date(e.get('expected_delivery_date')):
                left_bits.append(Paragraph(
                    _pdf_mixed_markup(f"{labels['delivery']}: {delivery_text(e)}", cjk_font),
                    overview_meta_style,
                ))
            left_bits.append(Paragraph(_pdf_mixed_markup(status_label, cjk_font), overview_status_style))
            timeline = MiniStageTimeline(
                overview_stage_index(e),
                overview_stage_dates(e),
                width=120 * mm,
                height=16 * mm,
                cancelled=(key == 'CANCELLED'),
            )
            # The status line is intentionally prominent.  Give rows with an
            # expected-delivery line a little more height so the red status never
            # gets squeezed against the divider below.
            overview_row_height = 25.0 * mm if _parse_db_date(e.get('expected_delivery_date')) else 23.0 * mm
            card = Table([[left_bits, timeline]], colWidths=[64 * mm, 120 * mm], rowHeights=[overview_row_height])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.white if ov_idx % 2 == 0 else sand_soft),
                ('LINEBEFORE', (0, 0), (0, 0), 1.8, accent),
                ('LINEBELOW', (0, 0), (-1, -1), .4, line),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (0, 0), 3.0 * mm),
                ('RIGHTPADDING', (0, 0), (0, 0), 2.0 * mm),
                ('LEFTPADDING', (1, 0), (1, 0), 0),
                ('RIGHTPADDING', (1, 0), (1, 0), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 1.8 * mm),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 1.8 * mm),
            ]))
            story.append(KeepTogether([card]))
            if ov_idx < len(entries) - 1:
                story.append(Spacer(1, 0.9 * mm))

        story.append(PageBreak())

    # ---------- Detailed order pages ----------
    for idx, entry in enumerate(entries):
        if idx > 0:
            story.append(PageBreak())

        order_display = entry.get('workflow_number') or entry.get('order_number') or '-'
        continued = text['continued'] if entry.get('continued') else ''
        key = status_key(entry)
        status_label = entry.get('status_label') or _report_status(entry.get('current_status'), language)

        # No boxy status pill: title + colored dot + text.
        status_markup = f'<font color="#FF2442">●</font> {_pdf_mixed_markup(status_label, cjk_font)}'
        order_head = Table([[
            [Paragraph(_pdf_mixed_markup(labels['detail'], cjk_font), eyebrow_style),
             Spacer(1, .8 * mm),
             Paragraph(f"{text['order_number']} {_pdf_mixed_markup(order_display, cjk_font)}{continued}", order_num_style)],
            Paragraph(status_markup, order_status_style),
        ]], colWidths=[125 * mm, 59 * mm])
        order_head.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0), ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(order_head)
        story.append(Spacer(1, 1.3 * mm))

        p_name = product_primary(entry)
        p_type = str(entry.get('production_type') or '').strip()
        p_code = str(entry.get('product_code') or '').strip()
        sub_bits = []
        if p_type and p_type != p_name:
            sub_bits.append(f"{labels['type']}: {p_type}")
        if p_code:
            sub_bits.append(f"{labels['code']}: {p_code}")
        sub_bits.append(f"{labels['order_date']}: {format_report_date(entry.get('order_date'), language) or '-'}")
        product_card = Table([[
            [Paragraph(_pdf_mixed_markup(labels['pattern'], cjk_font), hero_label_style),
             Paragraph(_pdf_mixed_markup(p_name, cjk_font), hero_value_style)],
            Paragraph(_pdf_mixed_markup('   /   '.join(sub_bits), cjk_font), hero_sub_style),
        ]], colWidths=[112 * mm, 72 * mm])
        product_card.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), sand_soft),
            ('LINEBEFORE', (0, 0), (0, 0), 1.8, accent),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (0, 0), 3.2 * mm),
            ('RIGHTPADDING', (0, 0), (0, 0), 3 * mm),
            ('LEFTPADDING', (1, 0), (1, 0), 3 * mm),
            ('RIGHTPADDING', (1, 0), (1, 0), 3 * mm),
            ('TOPPADDING', (0, 0), (-1, -1), 2.0 * mm),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2.0 * mm),
        ]))
        story.append(product_card)
        story.append(Spacer(1, 1.3 * mm))

        def key_card(label, value, value_color=None):
            val_style = key_value_style
            if value_color:
                val_style = ParagraphStyle(
                    f'kv{uuid.uuid4().hex[:6]}', parent=key_value_style, textColor=value_color
                )
            return [
                Paragraph(_pdf_mixed_markup(label, cjk_font), key_label_style),
                Paragraph(_pdf_mixed_markup(value or '-', cjk_font), val_style),
            ]

        qty = str(entry.get('quantity') or '-').strip() or '-'
        has_delivery = bool(_parse_db_date(entry.get('expected_delivery_date')))
        if has_delivery:
            card_items = [
                key_card(labels['delivery'], delivery_text(entry), accent_dark),
                key_card(labels['quantity'], qty),
                key_card(labels['current_status_card'], status_label, accent_dark),
            ]
            card_widths = [61.3 * mm, 61.3 * mm, 61.4 * mm]
        else:
            # 未确认交期不占版面，把空间留给数量和当前状态。
            status_card_style = ParagraphStyle(
                f'kvStatus{uuid.uuid4().hex[:6]}', parent=key_value_style,
                fontSize=12.0, leading=14.0, textColor=accent_dark,
            )
            card_items = [
                key_card(labels['quantity'], qty),
                [
                    Paragraph(_pdf_mixed_markup(labels['current_status_card'], cjk_font), key_label_style),
                    Paragraph(_pdf_mixed_markup(status_label or '-', cjk_font), status_card_style),
                ],
            ]
            card_widths = [92 * mm, 92 * mm]
        cards = Table([[*card_items]], colWidths=card_widths)
        cards.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            ('BOX', (0, 0), (-1, -1), .45, line),
            ('INNERGRID', (0, 0), (-1, -1), .35, line),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 3.6 * mm),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3.6 * mm),
            ('TOPPADDING', (0, 0), (-1, -1), 2.0 * mm),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2.0 * mm),
        ]))
        story.append(cards)
        story.append(Spacer(1, 2.0 * mm))

        progress_header = Table([[
            Paragraph(_pdf_mixed_markup(labels['current_progress'], cjk_font), section_style),
            Paragraph(_pdf_mixed_markup(status_label, cjk_font), small_style),
        ]], colWidths=[39 * mm, 145 * mm])
        progress_header.setStyle(TableStyle([
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'), ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0), ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(progress_header)
        story.append(Spacer(1, .5 * mm))
        story.append(StageTimeline(stage_index(key), width=184 * mm, height=13.5 * mm))
        story.append(Spacer(1, 1.4 * mm))

        shipping_records = entry.get('shipping_records') or []
        shipping_value = '<br/>'.join(_pdf_mixed_markup(x, cjk_font) for x in shipping_records) if shipping_records else _pdf_mixed_markup(text['no_shipping_history'], cjk_font)
        shipping = Table([[
            Paragraph(_pdf_mixed_markup(text['shipping_history'], cjk_font), section_style),
            Paragraph(shipping_value, body_style),
        ]], colWidths=[38 * mm, 146 * mm])
        shipping.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FFF7F8')),
            ('LINEBEFORE', (0, 0), (0, 0), 2.0, accent),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 3.4 * mm),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3.4 * mm),
            ('TOPPADDING', (0, 0), (-1, -1), 1.8 * mm),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1.8 * mm),
        ]))
        story.append(shipping)

        images = entry.get('images') or []
        if images:
            story.append(Spacer(1, 2.4 * mm))
            sections = _pdf_image_sections(entry, text)
            for section_idx, (section_label, section_images, combined) in enumerate(sections):
                if not section_images:
                    continue
                section_header = Table([[
                    Paragraph(_pdf_mixed_markup(section_label, cjk_font), section_style),
                    Paragraph(str(len(section_images)), small_style),
                ]], colWidths=[172 * mm, 12 * mm])
                section_header.setStyle(TableStyle([
                    ('LINEBELOW', (0, 0), (-1, -1), .45, line),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0), ('BOTTOMPADDING', (0, 0), (-1, -1), 1.4 * mm),
                ]))
                story.append(section_header)
                story.append(Spacer(1, 1.6 * mm))
                grid = _build_pdf_image_grid_two_columns(
                    section_images, RLImage, Paragraph, Table, TableStyle,
                    colors, mm, image_caption_style, image_source_style, cjk_font, text,
                    show_source=combined,
                )
                if grid is not None:
                    story.append(grid)
                if section_idx < len(sections) - 1:
                    story.append(Spacer(1, 2.7 * mm))

    def add_page_chrome(canvas, doc_obj):
        canvas.saveState()
        canvas.setStrokeColor(line)
        canvas.setLineWidth(.35)
        canvas.line(13 * mm, 10 * mm, A4[0] - 13 * mm, 10 * mm)
        footer_font = cjk_font if language == 'zh_cn' else 'Helvetica'
        canvas.setFont(footer_font, 7.0)
        canvas.setFillColor(muted)
        footer_left = customer_name or text['title']
        if language == 'es' and any(ord(ch) > 255 for ch in footer_left):
            footer_left = text['title']
        canvas.drawString(13 * mm, 6.2 * mm, footer_left[:52])
        canvas.drawRightString(A4[0] - 13 * mm, 6.2 * mm, f"{doc_obj.page:02d}")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_chrome, onLaterPages=add_page_chrome)
    return out.getvalue()

def _escape_pdf_text(value) -> str:
    from xml.sax.saxutils import escape
    return escape(str(value or ''))


def _pdf_mixed_markup(value, cjk_font: str = 'STSong-Light') -> str:
    """Keep Spanish/Latin in the normal PDF font while wrapping CJK runs in a CJK font."""
    from xml.sax.saxutils import escape
    text = str(value or '')
    if not text:
        return ''
    parts = []
    buf = []
    in_cjk = None

    def is_cjk(ch: str) -> bool:
        code = ord(ch)
        return (
            0x3400 <= code <= 0x4DBF or 0x4E00 <= code <= 0x9FFF or
            0xF900 <= code <= 0xFAFF or 0x3000 <= code <= 0x303F or
            0xFF00 <= code <= 0xFFEF
        )

    def flush():
        nonlocal buf, in_cjk
        if not buf:
            return
        chunk = escape(''.join(buf))
        parts.append(f'<font name="{cjk_font}">{chunk}</font>' if in_cjk else chunk)
        buf = []

    for ch in text:
        flag = is_cjk(ch)
        if in_cjk is None:
            in_cjk = flag
        elif flag != in_cjk:
            flush()
            in_cjk = flag
        buf.append(ch)
    flush()
    return ''.join(parts)


def _pdf_image_groups(entry: dict, text: dict) -> List[Tuple[str, List[dict]]]:
    """Return source groups respecting the user's selected priority."""
    images = entry.get('images') or []
    order_images = [x for x in images if x.get('source') == 'order']
    workflow_images = [x for x in images if x.get('source') == 'workflow']
    image_order = entry.get('image_order') or 'order_first'

    if image_order == 'workflow_first':
        groups = [(text['image_source_workflow'], workflow_images), (text['image_source_order'], order_images)]
    elif image_order == 'newest':
        def newest(group):
            return max((str(x.get('uploaded_at') or '') for x in group), default='')
        groups = [(text['image_source_order'], order_images), (text['image_source_workflow'], workflow_images)]
        groups = sorted(groups, key=lambda pair: newest(pair[1]), reverse=True)
    else:
        groups = [(text['image_source_order'], order_images), (text['image_source_workflow'], workflow_images)]
    return [(label, group) for label, group in groups if group]


def _pdf_image_sections(entry: dict, text: dict):
    """If each source has exactly one image, put both on the same visual row."""
    images = entry.get('images') or []
    order_images = [x for x in images if x.get('source') == 'order']
    workflow_images = [x for x in images if x.get('source') == 'workflow']
    if len(order_images) == 1 and len(workflow_images) == 1:
        image_order = entry.get('image_order') or 'order_first'
        if image_order == 'workflow_first':
            combined = workflow_images + order_images
        elif image_order == 'newest':
            combined = sorted(order_images + workflow_images, key=lambda x: str(x.get('uploaded_at') or ''), reverse=True)
        else:
            combined = order_images + workflow_images
        return [(text['images'], combined, True)]
    return [(label, group, False) for label, group in _pdf_image_groups(entry, text)]


def _build_pdf_image_grid_two_columns(images, RLImage, Paragraph, Table, TableStyle, colors, mm,
                                      caption_style, source_style, cjk_font, text, show_source=False):
    """Maximum two images per row. A third image always starts a new row."""
    if not images:
        return None
    total_w = 178 * mm
    gap = 5 * mm
    cell_w = (total_w - gap) / 2
    max_img_h = (76 if show_source else 68) * mm
    rows = []
    current = []

    for img in images:
        try:
            width_pt, height_pt = _fit_size(
                img['width'], img['height'],
                max_w=cell_w - 7 * mm,
                max_h=max_img_h,
            )
            pic = RLImage(BytesIO(img['bytes']), width=width_pt, height=height_pt)
            pic.hAlign = 'CENTER'
            source_label = text['image_source_order'] if img.get('source') == 'order' else text['image_source_workflow']
            filename = str(img.get('display_name') or '').strip()
            card_rows = []
            if show_source:
                card_rows.append([Paragraph(_pdf_mixed_markup(source_label, cjk_font), source_style)])
            card_rows.append([pic])
            if filename:
                card_rows.append([Paragraph(_pdf_mixed_markup(filename, cjk_font), caption_style)])
            card = Table(card_rows, colWidths=[cell_w - 2 * mm])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                ('BOX', (0, 0), (-1, -1), .45, colors.HexColor('#E0E6EC')),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 2.2 * mm),
                ('RIGHTPADDING', (0, 0), (-1, -1), 2.2 * mm),
                ('TOPPADDING', (0, 0), (-1, -1), 2.2 * mm),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2.2 * mm),
            ]))
            current.append(card)
            if len(current) == 2:
                rows.append(current)
                current = []
        except Exception:
            continue

    if current:
        current.append('')
        rows.append(current)
    if not rows:
        return None

    grid = Table(rows, colWidths=[cell_w, cell_w], hAlign='LEFT', splitByRow=1)
    grid.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), gap / 2),
        ('TOPPADDING', (0, 0), (-1, -1), 1.1 * mm),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 1.8 * mm),
    ]))
    return grid

def _fit_size(width: int, height: int, max_w: float, max_h: float) -> Tuple[float, float]:
    if not width or not height:
        return max_w, max_h
    scale = min(max_w / float(width), max_h / float(height), 1.0)
    return width * scale, height * scale


def _render_docx(entries: List[dict], language: str) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    text = REPORT_TEXT[language]
    body_font = 'Microsoft YaHei' if language == 'zh_cn' else 'Arial'

    normal = doc.styles['Normal']
    normal.font.name = body_font
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text['title'])
    run.bold = True
    run.font.size = Pt(18)
    _set_run_font(run, body_font)

    customers = _distinct_customers(entries)
    if len(customers) == 1:
        p = doc.add_paragraph()
        p.add_run(f"{text['customer']}: ").bold = True
        p.add_run(customers[0])
    p = doc.add_paragraph()
    p.add_run(f"{text['order_count']}: ").bold = True
    p.add_run(str(len(entries)))
    p.add_run(f"    {text['generated_at']}: ").bold = True
    p.add_run(format_report_datetime(datetime.now(), language))

    for idx, entry in enumerate(entries):
        if idx > 0:
            doc.add_page_break()
        order_display = entry.get('workflow_number') or entry.get('order_number') or '-'
        continued = text['continued'] if entry.get('continued') else ''
        heading = doc.add_paragraph()
        run = heading.add_run(f"{text['order_number']}: {order_display}{continued}")
        run.bold = True
        run.font.size = Pt(13)
        _set_run_font(run, body_font)

        fields = _report_fields(entry, language)
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = 'Table Grid'
        for row, (label, value) in zip(table.rows, fields):
            row.cells[0].text = label
            row.cells[1].text = str(value or '-')
            for r in row.cells[0].paragraphs[0].runs:
                r.bold = True
                _set_run_font(r, body_font)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        _set_run_font(r, body_font)

        p = doc.add_paragraph()
        r = p.add_run(text['shipping_history'] + ': ')
        r.bold = True
        _set_run_font(r, body_font)
        shipping = entry.get('shipping_records') or []
        if shipping:
            p.add_run(shipping[0])
            for item in shipping[1:]:
                p.add_run('\n' + item)
        else:
            p.add_run(text['no_shipping_history'])
        for r in p.runs:
            _set_run_font(r, body_font)

        images = entry.get('images') or []
        if images:
            p = doc.add_paragraph()
            r = p.add_run(text['images'])
            r.bold = True
            _set_run_font(r, body_font)
            for img in images:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    w_in, h_in = _fit_size(img['width'], img['height'], max_w=6.6, max_h=4.6)
                    p.add_run().add_picture(BytesIO(img['bytes']), width=Inches(w_in), height=Inches(h_in))
                except Exception:
                    continue

    # Set all remaining run fonts, including table text created by python-docx.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, body_font)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, body_font)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _set_run_font(run, font_name: str) -> None:
    from docx.oxml.ns import qn
    run.font.name = font_name
    # Español 报告内的产品名称可能仍是中文；East Asia 永远指定可显示中文的字体。
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _render_excel(entries: List[dict], language: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    text = REPORT_TEXT[language]
    wb = Workbook()
    ws = wb.active
    ws.title = text['report_sheet'][:31]

    # Keep the two image sources visible on the main Excel sheet. Previously
    # there was only one generic image column, so "both" always looked like it
    # contained only the first source (normally the supervisor reference image).
    headers = [
        text['order_number'], text['order_date'], text['customer'], text['product_name'],
        text['product_type'], text['product_code'], text['quantity'], text['status'],
        text['delivery_date'], text['shipping_history'],
        text['image_source_order'], text['image_source_workflow']
    ]
    header_fill = PatternFill('solid', fgColor='DCEBFA')
    header_font = Font(bold=True, color='1F2937')
    thin = Side(style='thin', color='D1D5DB')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col, header in enumerate(headers, 1):
        cell = ws.cell(1, col, header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    for row_idx, entry in enumerate(entries, 2):
        shipping = '\n'.join(entry.get('shipping_records') or []) or text['no_shipping_history']
        values = [
            entry.get('workflow_number') or entry.get('order_number') or '',
            format_report_date(entry.get('order_date'), language),
            entry.get('customer_name') or '', entry.get('product_name') or '',
            entry.get('production_type') or '', entry.get('product_code') or '',
            entry.get('quantity') or '', entry.get('status_label') or '',
            format_report_date(entry.get('expected_delivery_date'), language), shipping, '', ''
        ]
        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row_idx, col_idx, value)
            cell.border = border
            cell.alignment = Alignment(vertical='top', wrap_text=True)

        images = entry.get('images') or []
        # Main sheet: always show one representative image from EACH selected source.
        # Do not derive these two cells from the combined representative list, because
        # that list may intentionally contain only one image total.
        source_reps = entry.get('excel_source_images') or {}
        order_img = source_reps.get('order') or next((img for img in images if img.get('source') == 'order'), None)
        workflow_img = source_reps.get('workflow') or next((img for img in images if img.get('source') == 'workflow'), None)
        for img, col_letter in ((order_img, 'K'), (workflow_img, 'L')):
            if not img:
                continue
            try:
                xl_img = XLImage(BytesIO(img['bytes']))
                max_w, max_h = 150, 88
                scale = min(max_w / xl_img.width, max_h / xl_img.height, 1.0)
                xl_img.width *= scale
                xl_img.height *= scale
                ws.add_image(xl_img, f'{col_letter}{row_idx}')
                ws.row_dimensions[row_idx].height = max(ws.row_dimensions[row_idx].height or 15, 70)
            except Exception:
                pass

    widths = [18, 15, 24, 22, 20, 18, 12, 24, 17, 38, 24, 24]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = 'A2'
    ws.auto_filter.ref = f'A1:L{max(1, len(entries) + 1)}'
    ws.row_dimensions[1].height = 24

    all_images = [(e, img) for e in entries for img in (e.get('images') or [])]
    # With "all images", every selected image is preserved in a dedicated sheet.
    # With representative image mode, this sheet contains only those representative images.
    if all_images:
        iw = wb.create_sheet(text['image_sheet'][:31])
        image_headers = [text['order_number'], text['image_source'], text['image_name'], text['image']]
        for col, header in enumerate(image_headers, 1):
            cell = iw.cell(1, col, header)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border
            cell.alignment = Alignment(horizontal='center')
        r = 2
        for entry, img in all_images:
            order_display = entry.get('workflow_number') or entry.get('order_number') or ''
            source_label = text['image_source_order'] if img.get('source') == 'order' else text['image_source_workflow']
            iw.cell(r, 1, order_display)
            iw.cell(r, 2, source_label)
            iw.cell(r, 3, img.get('display_name') or os.path.basename(img.get('path') or ''))
            for c in range(1, 5):
                iw.cell(r, c).border = border
                iw.cell(r, c).alignment = Alignment(vertical='top', wrap_text=True)
            try:
                xl_img = XLImage(BytesIO(img['bytes']))
                max_w, max_h = 220, 130
                scale = min(max_w / xl_img.width, max_h / xl_img.height, 1.0)
                xl_img.width *= scale
                xl_img.height *= scale
                iw.add_image(xl_img, f'D{r}')
                iw.row_dimensions[r].height = 100
            except Exception:
                pass
            r += 1
        iw.column_dimensions['A'].width = 20
        iw.column_dimensions['B'].width = 24
        iw.column_dimensions['C'].width = 34
        iw.column_dimensions['D'].width = 34
        iw.freeze_panes = 'A2'

    out = BytesIO()
    wb.save(out)
    return out.getvalue()


def _report_fields(entry: dict, language: str) -> List[Tuple[str, str]]:
    text = REPORT_TEXT[language]
    return [
        (text['order_date'], format_report_date(entry.get('order_date'), language)),
        (text['customer'], entry.get('customer_name') or ''),
        (text['product_name'], entry.get('product_name') or ''),
        (text['product_type'], entry.get('production_type') or ''),
        (text['product_code'], entry.get('product_code') or ''),
        (text['quantity'], str(entry.get('quantity') or '')),
        (text['status'], entry.get('status_label') or _report_status(entry.get('current_status'), language)),
        (text['delivery_date'], format_report_date(entry.get('expected_delivery_date'), language)),
    ]


def _distinct_customers(entries: Sequence[dict]) -> List[str]:
    out = []
    for e in entries:
        name = str(e.get('customer_name') or '').strip()
        if name and name not in out:
            out.append(name)
    return out


def ensure_report_cache() -> str:
    os.makedirs(CUSTOMER_REPORT_CACHE_DIR, exist_ok=True)
    cleanup_report_cache()
    return CUSTOMER_REPORT_CACHE_DIR


def cleanup_report_cache() -> None:
    if not os.path.isdir(CUSTOMER_REPORT_CACHE_DIR):
        return
    cutoff = datetime.now() - timedelta(hours=CUSTOMER_REPORT_CACHE_HOURS)
    for name in os.listdir(CUSTOMER_REPORT_CACHE_DIR):
        path = os.path.join(CUSTOMER_REPORT_CACHE_DIR, name)
        try:
            if os.path.isfile(path) and datetime.fromtimestamp(os.path.getmtime(path)) < cutoff:
                os.remove(path)
        except OSError:
            pass


def cache_report_file(filename: str, mimetype: str, data: bytes) -> dict:
    ensure_report_cache()
    file_id = uuid.uuid4().hex
    data_path = os.path.join(CUSTOMER_REPORT_CACHE_DIR, f'{file_id}.bin')
    meta_path = os.path.join(CUSTOMER_REPORT_CACHE_DIR, f'{file_id}.json')
    with open(data_path, 'wb') as f:
        f.write(data)
    meta = {'filename': filename, 'mimetype': mimetype, 'size': len(data), 'created_at': datetime.now().isoformat()}
    with open(meta_path, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False)
    return {'id': file_id, **meta}


def get_cached_report(file_id: str) -> Optional[dict]:
    if not re.fullmatch(r'[0-9a-f]{32}', str(file_id or '')):
        return None
    data_path = os.path.join(CUSTOMER_REPORT_CACHE_DIR, f'{file_id}.bin')
    meta_path = os.path.join(CUSTOMER_REPORT_CACHE_DIR, f'{file_id}.json')
    if not (os.path.isfile(data_path) and os.path.isfile(meta_path)):
        return None
    cutoff = datetime.now() - timedelta(hours=CUSTOMER_REPORT_CACHE_HOURS)
    try:
        if datetime.fromtimestamp(os.path.getmtime(data_path)) < cutoff:
            for path in (data_path, meta_path):
                try:
                    os.remove(path)
                except OSError:
                    pass
            return None
        with open(meta_path, 'r', encoding='utf-8') as f:
            meta = json.load(f)
    except Exception:
        return None
    return {'path': data_path, **meta}
